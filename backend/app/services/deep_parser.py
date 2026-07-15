
"""
文档深度解析 - OCR 图片 + 表格提取 + 跨页表格合并
"""
import os
import re
import csv
import io
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from PIL import Image


# ============================================================
# OCR
# ============================================================

def ocr_image(image_path: str, lang: str = "chi_sim+eng") -> str:
    """OCR 识别图片文字（优先 PaddleOCR，降级 tesseract）"""
    try:
        # PaddleOCR（识别率更高）
        from paddleocr import PaddleOCR
        _paddle = PaddleOCR(lang='ch')
        result = _paddle.ocr(image_path)
        if result and result[0]:
            texts = [line[1][0] for line in result[0] if line and len(line) > 1]
            return "\n".join(texts).strip()
    except Exception:
        pass

    try:
        # 降级 tesseract
        import pytesseract
        img = Image.open(image_path)
        # 提升识别率：放大 + 灰度化
        w, h = img.size
        if w < 500:
            img = img.resize((w * 2, h * 2), Image.LANCZOS)
        if img.mode != "L":
            img = img.convert("L")
        return pytesseract.image_to_string(img, lang=lang).strip()
    except Exception:
        return ""


# ============================================================
# 图片提取
# ============================================================

def extract_images_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """从 PDF 提取嵌入图片"""
    images = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                for img in page.images:
                    images.append({
                        "page": page_num + 1,
                        "bbox": (img["x0"], img["top"], img["x1"], img["bottom"]),
                        "width": img["width"],
                        "height": img["height"],
                    })
    except Exception:
        pass
    return images


# ============================================================
# 表格提取 + 跨页合并
# ============================================================

def _table_header_hash(rows: List[List[str]]) -> str:
    """计算表头哈希（用于判断相邻页表格是否同一个）"""
    if not rows or not rows[0]:
        return ""
    return "|".join(str(c).strip() for c in rows[0])


def _merge_similar_tables(tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    合并跨页表格

    策略: 如果相邻两个表格列数相同且表头相似（或后一页表格无表头），则合并
    """
    if len(tables) <= 1:
        return tables

    merged = []
    current = None

    for t in tables:
        if current is None:
            current = t
            continue

        # 判断是否该合并
        same_cols = (current["cols"] == t["cols"])
        curr_header = _table_header_hash(current.get("data", []))
        next_header = _table_header_hash(t.get("data", []))

        # 列数相同 + (表头相同 或 后一表格无有效表头)
        if same_cols and (curr_header == next_header or not next_header or len(t.get("data", [])) <= 1):
            # 去掉后一表格的重复表头行
            next_rows = t["data"]
            if curr_header == next_header and len(next_rows) > 1:
                next_rows = next_rows[1:]  # 跳过重复表头
            elif len(next_rows) > 0:
                # 检查第一行是否像表头（与当前表头相似）
                first_row_hash = "|".join(str(c).strip() for c in next_rows[0])
                if _similarity(curr_header, first_row_hash) > 0.7:
                    next_rows = next_rows[1:]

            current["data"].extend(next_rows)
            current["rows"] = len(current["data"])
            current["end_page"] = t.get("page", current.get("page", 1))
            current["merged"] = True
        else:
            merged.append(current)
            current = t

    if current:
        merged.append(current)
    return merged


def _similarity(a: str, b: str) -> float:
    """简单字符串相似度"""
    if not a or not b:
        return 0.0
    set_a, set_b = set(a.split("|")), set(b.split("|"))
    if not set_a:
        return 0.0
    return len(set_a & set_b) / len(set_a)


def extract_tables_from_pdf(file_path: str, merge_cross_page: bool = True) -> List[Dict[str, Any]]:
    """从 PDF 提取表格（支持跨页合并）"""
    tables = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    if table:
                        cleaned = [[str(c).strip() if c else "" for c in row]
                                   for row in table if any(c for c in row)]
                        if cleaned:
                            tables.append({
                                "page": page_num + 1,
                                "table_index": table_idx,
                                "rows": len(cleaned),
                                "cols": len(cleaned[0]) if cleaned else 0,
                                "data": cleaned,
                            })
    except Exception:
        pass

    if merge_cross_page:
        tables = _merge_similar_tables(tables)
    return tables


def extract_tables_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """从 Word 提取表格"""
    tables = []
    try:
        from docx import Document
        doc = Document(file_path)
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                tables.append({
                    "table_index": table_idx,
                    "rows": len(rows),
                    "cols": len(rows[0]) if rows else 0,
                    "data": rows,
                })
    except Exception:
        pass
    return tables


def tables_to_text(tables: List[Dict[str, Any]]) -> str:
    """将表格数据转为可检索的 Markdown 表格文本"""
    parts = []
    for t in tables:
        page_info = f"第{t.get('page', '?')}页"
        end_info = f"-{t['end_page']}页" if t.get("end_page") and t["end_page"] != t.get("page") else ""
        merged_tag = " [跨页合并]" if t.get("merged") else ""
        parts.append(f"\n[表格{merged_tag} {page_info}{end_info} — {t['rows']}行 × {t['cols']}列]")

        # 表头
        if t["data"]:
            parts.append("| " + " | ".join(t["data"][0]) + " |")
            parts.append("|" + "|".join(["---"] * t["cols"]) + "|")
            # 数据行（最多展示 50 行避免过长）
            for row in t["data"][1:51]:
                padded = row + [""] * (t["cols"] - len(row))
                parts.append("| " + " | ".join(padded[:t["cols"]]) + " |")
            if len(t["data"]) > 51:
                parts.append(f"... (共 {len(t['data']) - 1} 行数据)")
    return "\n".join(parts)


# ============================================================
# 深度解析入口
# ============================================================

def deep_parse_pdf(file_path: str) -> dict:
    """PDF 深度解析：文本 + OCR 图片 + 跨页表格合并"""
    import pdfplumber
    text_parts = []
    ocr_texts = []
    images_found = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 文本
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # 图片 OCR
                for img in page.images:
                    images_found.append({
                        "page": page_num + 1,
                        "bbox": (img["x0"], img["top"], img["x1"], img["bottom"]),
                    })
                    try:
                        bbox = (int(img["x0"]), int(img["top"]), int(img["x1"]), int(img["bottom"]))
                        cropped = page.within_bbox(bbox).to_image()
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                            cropped.save(tmp.name)
                            ocr_text = ocr_image(tmp.name)
                            if ocr_text:
                                ocr_texts.append(f"[图片OCR 第{page_num + 1}页] {ocr_text}")
                            os.unlink(tmp.name)
                    except Exception:
                        pass
    except Exception:
        pass

    # 跨页表格
    tables_found = extract_tables_from_pdf(file_path, merge_cross_page=True)

    full_text = "\n\n".join(text_parts)
    if tables_found:
        full_text += "\n\n" + tables_to_text(tables_found)
    if ocr_texts:
        full_text += "\n\n" + "\n".join(ocr_texts)

    return {
        "text": full_text,
        "tables_count": len(tables_found),
        "images_found": len(images_found),
        "images_ocr_count": len(ocr_texts),
    }


def deep_parse_docx(file_path: str) -> dict:
    """Word 深度解析：文本 + 表格"""
    from docx import Document
    doc = Document(file_path)
    text_parts = []
    tables_found = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables_found.append({
                "table_index": table_idx,
                "rows": len(rows),
                "cols": len(rows[0]),
                "data": rows,
            })

    full_text = "\n\n".join(text_parts)
    if tables_found:
        full_text += "\n\n" + tables_to_text(tables_found)

    return {
        "text": full_text,
        "tables_count": len(tables_found),
        "images_found": 0,
        "images_ocr_count": 0,
    }


def deep_parse(file_path: str) -> dict:
    """统一的深度解析入口"""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return deep_parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return deep_parse_docx(file_path)
    else:
        from app.services.document_parser import extract_text
        return {
            "text": extract_text(file_path),
            "tables_count": 0,
            "images_found": 0,
            "images_ocr_count": 0,
        }
